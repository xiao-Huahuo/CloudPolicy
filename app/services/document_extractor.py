from pathlib import Path

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None


def _load_openpyxl():
    """
    延迟导入 openpyxl，避免在应用启动阶段因 numpy/openpyxl 环境问题导致崩溃。
    """
    try:
        import openpyxl as _openpyxl
        return _openpyxl, None
    except Exception as e:
        return None, e


def _load_xlrd():
    try:
        import xlrd as _xlrd
        return _xlrd, None
    except Exception as e:
        return None, e


def extract_text_from_docx(file_path: Path) -> str:
    """提取 DOCX 文件中的文字（使用 python-docx）。"""
    if docx is None:
        return "警告：由于未安装 python-docx 库，无法提取 Word 文件内容。"

    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 尝试提取表格中的文字
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(cell.text.strip())
                if row_data:
                    full_text.append(" | ".join(row_data))

        return "\n".join(full_text)
    except Exception as e:
        return f"DOCX 解析失败: {str(e)}"


def extract_text_from_excel(file_path: Path) -> str:
    """提取 Excel (xlsx/xls) 文件中的文字。"""
    ext = file_path.suffix.lower()
    full_text = []

    try:
        if ext == ".xlsx":
            openpyxl, openpyxl_err = _load_openpyxl()
            if openpyxl is None:
                return f"Excel 解析失败: 无法加载 openpyxl ({openpyxl_err})"

            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet in wb.worksheets:
                full_text.append(f"--- 表格: {sheet.title} ---")
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))

        elif ext == ".xls":
            xlrd, xlrd_err = _load_xlrd()
            if xlrd is None:
                return f"Excel 解析失败: 无法加载 xlrd ({xlrd_err})"

            wb = xlrd.open_workbook(file_path)
            for sheet in wb.sheets():
                full_text.append(f"--- 表格: {sheet.name} ---")
                for row_idx in range(sheet.nrows):
                    row_values = sheet.row_values(row_idx)
                    row_data = [str(val).strip() for val in row_values if val and str(val).strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))
        else:
            return f"Excel 解析失败: 不支持的扩展名 {ext}"

        return "\n".join(full_text)
    except Exception as e:
        return f"Excel 解析失败: {str(e)}"
